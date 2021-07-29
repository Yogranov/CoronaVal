import xlrd 
from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Cm, Inches


def set_col_widths(table):
    widths = (Cm(0.5), Inches(1.5), Inches(1.5), Inches(1), Inches(1.5), Inches(0.1))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width



wb = xlrd.open_workbook("input.xlsx") 
sheet = wb.sheet_by_index(0) 
  
pathList = {}

for x in range(sheet.nrows):
    row = sheet.row(x)

    if row[1].value.isdigit() != True:
        continue

    if row[1].value in pathList:
        pathList[row[1].value].append(row)
    else:
        pathList[row[1].value] = [row]


document = Document()

style = document.styles['Heading2']
font = style.font
font.size = Pt(24)
font.color.rgb = RGBColor(0, 0, 0)

valPhone = '050-000-0000' # Volunteers Phone Number
count = 1
for x in pathList:
    total = 0
    numOfRows = 1#len(pathList[x])
    
    title = document.add_heading('חבילה מספר ' + str(count), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=numOfRows, cols=7)
    table.style = 'TableGrid'

    set_col_widths(table)


    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'כמות'
    hdr_cells[1].text = 'טלפון נוסף'
    hdr_cells[2].text = 'טלפון'
    hdr_cells[3].text = 'דירה וקומה'
    hdr_cells[4].text = 'כתובת'
    hdr_cells[5].text = 'שם'
    hdr_cells[6].text = 'מס עצירה'

    for i in pathList[x]:
        quantity = int(i[4].value)
        phone = '0' + i[7].value[4:]
        anotherPhone = i[6].value
        appartment = i[8].value
        address = i[9].value
        name = i[10].value
        stopNum = i[11].value

        total += quantity

        row_cells = table.add_row().cells
        row_cells[0].text = str(quantity)
        row_cells[1].text = str(anotherPhone)
        row_cells[2].text = str(phone)
        row_cells[3].text = str(appartment)
        row_cells[4].text = str(address)
        row_cells[5].text = str(name)
        row_cells[6].text = str(int(stopNum))

    document.add_paragraph('    ').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('סהכ חבילות: ' + str(total)).alignment = WD_ALIGN_PARAGRAPH.RIGHT


    document.add_paragraph('מספר נתיב: ' + str(x)).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('    ').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('טלפון מתנדבים: ' + valPhone, style='Heading2').alignment = WD_ALIGN_PARAGRAPH.CENTER

    count +=1
    document.add_page_break()


document.save('output.docx')